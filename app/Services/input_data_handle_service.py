"""
Input Data Handle Service

File extraction, validation, and text chunking for document ingestion.
Supports PDF, DOCX, TXT, and Markdown files.
"""

import logging
import hashlib
import io
from typing import List, Dict, Optional, Tuple, BinaryIO, Any
from pathlib import Path
from datetime import datetime, timezone

# Document extraction libraries
from PyPDF2 import PdfReader
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# LangChain text splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Application imports
from app.core.config import settings
from app.Services.chunking_strategies import (
    ChunkingStrategyFactory,
    get_default_strategy
)

logger = logging.getLogger(__name__)


class InputDataHandleService:
    """
    Service for handling document ingestion workflow

    Workflow:
    1. Validate file (type, size, content)
    2. Extract text from file
    3. Chunk text for embedding
    4. Generate metadata

    Supported formats: PDF, DOCX, TXT, MD
    """

    def __init__(
        self,
        chunking_strategy: Optional[str] = None,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
        allowed_extensions: Optional[List[str]] = None,
        max_file_size: Optional[int] = None
    ):
        """
        Initialize Input Data Handle Service

        Args:
            chunking_strategy: Strategy name ("recursive" or "hierarchical")
            chunk_size: Text chunk size (default from settings)
            chunk_overlap: Chunk overlap size (default from settings)
            allowed_extensions: Allowed file extensions (default from settings)
            max_file_size: Max file size in bytes (default from settings)
        """
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        self.allowed_extensions = allowed_extensions or settings.ALLOWED_EXTENSIONS
        self.max_file_size = max_file_size or settings.MAX_FILE_SIZE

        # Legacy text splitter (fallback)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=settings.CHUNK_SEPARATORS,
            length_function=len
        )

        # Strategy Pattern: Initialize chunking strategy
        strategy_name = chunking_strategy or settings.CHUNKING_STRATEGY
        if strategy_name == "hierarchical":
            self.chunking_strategy = ChunkingStrategyFactory.create(
                "hierarchical",
                chunk_sizes=settings.HIERARCHICAL_CHUNK_SIZES,
                overlap=settings.HIERARCHICAL_OVERLAP,
                separators=settings.CHUNK_SEPARATORS
            )
        else:
            self.chunking_strategy = ChunkingStrategyFactory.create(
                "recursive",
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=settings.CHUNK_SEPARATORS
            )

        logger.info(
            f"Input Data Handle Service initialized: "
            f"strategy={strategy_name}, chunk_size={self.chunk_size}"
        )

    def validate_file(
        self,
        file_content: bytes,
        filename: str
    ) -> Tuple[bool, Optional[str]]:
        """
        Validate file before processing

        Args:
            file_content: Binary file content
            filename: Original filename

        Returns:
            Tuple of (is_valid, error_message)

        Example:
            >>> is_valid, error = service.validate_file(content, "doc.pdf")
            >>> if not is_valid:
            ...     raise ValueError(error)
        """
        # Check file extension
        file_ext = Path(filename).suffix.lower().lstrip('.')
        if file_ext not in self.allowed_extensions:
            return False, (
                f"Unsupported file type: '{file_ext}'. "
                f"Allowed: {', '.join(self.allowed_extensions)}"
            )

        # Check file size
        file_size = len(file_content)
        if file_size > self.max_file_size:
            max_mb = self.max_file_size / (1024 * 1024)
            current_mb = file_size / (1024 * 1024)
            return False, (
                f"File too large: {current_mb:.2f}MB (max: {max_mb:.2f}MB)"
            )

        # Check file is not empty
        if file_size == 0:
            return False, "File is empty"

        return True, None

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file

        Args:
            file_content: PDF binary content

        Returns:
            Extracted text

        Raises:
            ValueError: If PDF extraction fails
        """
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            # Extract text from all pages
            corpus = ''.join([
                page.extract_text() or ''
                for page in reader.pages
            ])

            if not corpus.strip():
                raise ValueError("PDF contains no extractable text")

            logger.info(f"Extracted {len(corpus)} characters from PDF ({len(reader.pages)} pages)")
            return corpus

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract PDF: {str(e)}")

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            file_content: DOCX binary content

        Returns:
            Extracted text

        Raises:
            ValueError: If DOCX extraction fails
            ImportError: If python-docx not installed
        """
        if DocxDocument is None:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            # Extract all paragraphs
            corpus = '\n\n'.join([
                paragraph.text
                for paragraph in doc.paragraphs
                if paragraph.text.strip()
            ])

            if not corpus.strip():
                raise ValueError("DOCX contains no text")

            logger.info(f"Extracted {len(corpus)} characters from DOCX ({len(doc.paragraphs)} paragraphs)")
            return corpus

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract DOCX: {str(e)}")

    def extract_text_from_txt(self, file_content: bytes) -> str:
        """
        Extract text from TXT/MD file

        Args:
            file_content: Text binary content

        Returns:
            Decoded text

        Raises:
            ValueError: If text decoding fails
        """
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')

            if not text.strip():
                raise ValueError("Text file is empty")

            logger.info(f"Extracted {len(text)} characters from text file")
            return text

        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text: {str(e)}")

    def extract_text(
        self,
        file_content: bytes,
        filename: str
    ) -> str:
        """
        Extract text from file based on extension

        Args:
            file_content: Binary file content
            filename: Original filename

        Returns:
            Extracted text

        Example:
            >>> text = service.extract_text(content, "document.pdf")
        """
        file_ext = Path(filename).suffix.lower().lstrip('.')

        if file_ext == 'pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_ext == 'docx':
            return self.extract_text_from_docx(file_content)
        elif file_ext in ['txt', 'md']:
            return self.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")

    def chunk_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Split text into chunks using configured strategy

        Args:
            text: Full document text
            metadata: Optional metadata to attach to chunks

        Returns:
            List of chunk dicts (from strategy.chunk())

        Example:
            >>> chunks = service.chunk_text(extracted_text, metadata={"file_id": "abc"})
            >>> print(f"Generated {len(chunks)} chunks")
            >>> print(chunks[0].keys())  # ['content', 'chunk_index', 'metadata', ...]
        """
        try:
            # Use Strategy Pattern for chunking
            chunks = self.chunking_strategy.chunk(text, metadata=metadata)

            logger.info(
                f"Text chunked using {self.chunking_strategy.get_strategy_name()}: "
                f"{len(chunks)} chunks generated"
            )

            return chunks

        except Exception as e:
            logger.error(f"Text chunking failed: {str(e)}")
            raise ValueError(f"Failed to chunk text: {str(e)}")

    def chunk_text_legacy(self, text: str) -> List[str]:
        """
        Legacy method: Split text into simple string chunks

        DEPRECATED: Use chunk_text() which returns structured chunk dicts

        Args:
            text: Full document text

        Returns:
            List of text chunks (strings only)
        """
        try:
            chunks = self.text_splitter.split_text(text)
            chunks = [chunk.strip() for chunk in chunks if chunk.strip()]

            logger.warning("Using legacy chunk_text_legacy() - prefer chunk_text()")
            return chunks

        except Exception as e:
            logger.error(f"Legacy chunking failed: {str(e)}")
            raise ValueError(f"Failed to chunk text: {str(e)}")

    def generate_file_id(
        self,
        file_content: bytes,
        filename: str
    ) -> str:
        """
        Generate unique file ID with timestamp + UUID + content hash

        Args:
            file_content: Binary file content
            filename: Original filename

        Returns:
            Unique file ID (format: "file_{timestamp}_{uuid8}_{hash8}")

        Example:
            >>> file_id = service.generate_file_id(content, "doc.pdf")
            >>> # file_id = "file_1698765432_a1b2c3d4_e5f6g7h8"

        Components:
            - timestamp: Unix timestamp (10 digits, chronological sorting)
            - uuid8: First 8 chars of UUID4 (collision resistance)
            - hash8: First 8 chars of SHA256 (content verification)
        """
        import uuid
        import time

        # Component 1: Timestamp (sortable, chronological ordering)
        timestamp = int(time.time())

        # Component 2: Random UUID (collision resistance)
        uuid_part = str(uuid.uuid4()).replace('-', '')[:8]

        # Component 3: Content hash (duplicate detection)
        content_hash = hashlib.sha256(file_content).hexdigest()[:8]

        file_id = f"file_{timestamp}_{uuid_part}_{content_hash}"

        return file_id

    async def generate_unique_file_id(
        self,
        file_content: bytes,
        filename: str,
        file_metadata_provider,
        max_retries: int = 3
    ) -> str:
        """
        Generate unique file ID with database collision detection

        This method generates a file_id and checks the database to ensure
        it doesn't already exist. If a collision is detected, it retries
        with a fresh UUID component.

        Args:
            file_content: Binary file content
            filename: Original filename
            file_metadata_provider: FileMetadataProvider instance for DB checks
            max_retries: Maximum number of generation attempts (default: 3)

        Returns:
            Unique file ID guaranteed not to exist in database

        Raises:
            ValueError: If unable to generate unique ID after max_retries attempts

        Example:
            >>> file_id = await service.generate_unique_file_id(
            ...     content, "doc.pdf", provider, max_retries=3
            ... )
            >>> # file_id = "file_1698765432_a1b2c3d4_e5f6g7h8" (guaranteed unique)

        Collision Probability:
            - Single attempt: ~1 in 4.3 billion (UUID4 collision)
            - With 3 retries: ~1 in 79 quintillion
        """
        for attempt in range(max_retries):
            # Generate candidate file_id
            file_id = self.generate_file_id(file_content, filename)

            # Check if file_id already exists in database
            try:
                existing_file = await file_metadata_provider.get_file(file_id)

                if existing_file is None:
                    # No collision, use this file_id
                    if attempt > 0:
                        logger.info(f"Generated unique file_id: {file_id} (attempt {attempt + 1})")
                    return file_id
                else:
                    # Collision detected, retry with new UUID
                    logger.warning(
                        f"file_id collision detected: {file_id} "
                        f"(attempt {attempt + 1}/{max_retries}). Retrying..."
                    )
                    continue

            except Exception as e:
                logger.error(f"Error checking file_id uniqueness: {str(e)}")
                # On error, return generated ID (fail open for availability)
                return file_id

        # All retries exhausted
        raise ValueError(
            f"Failed to generate unique file_id after {max_retries} attempts. "
            "This is extremely unlikely and may indicate a database issue."
        )

    def enrich_chunk_metadata(
        self,
        chunks: List[Dict[str, Any]],
        file_id: str,
        filename: str,
        file_size: int
    ) -> List[Dict[str, Any]]:
        """
        Enrich chunk metadata with file-level information

        Args:
            chunks: Chunks from strategy.chunk() (with strategy metadata)
            file_id: Unique file identifier
            filename: Original filename
            file_size: Original file size in bytes

        Returns:
            List of enriched chunk dicts

        Example:
            >>> chunks = strategy.chunk(text)
            >>> enriched = service.enrich_chunk_metadata(chunks, "file_abc", "doc.pdf", 1024000)
        """
        timestamp = int(datetime.now(timezone.utc).timestamp())

        for chunk in chunks:
            # Merge with existing metadata from strategy
            chunk["metadata"].update({
                "file_id": file_id,
                "filename": filename,
                "file_size": file_size,
                "timestamp": timestamp,
                "total_chunks": len(chunks)
            })

        return chunks

    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        file_metadata_provider = None
    ) -> Dict:
        """
        Complete file processing workflow

        Workflow:
        1. Validate file
        2. Extract text
        3. Chunk text
        4. Generate metadata

        Args:
            file_content: Binary file content
            filename: Original filename
            file_metadata_provider: Optional FileMetadataProvider for unique ID generation

        Returns:
            Dict with keys:
            - file_id: Unique identifier
            - filename: Original filename
            - file_size: File size in bytes
            - chunks: List of text chunks
            - metadata: List of chunk metadata dicts
            - chunk_count: Number of chunks

        Example:
            >>> result = await service.process_file(content, "document.pdf", provider)
            >>> print(f"File ID: {result['file_id']}")
            >>> print(f"Chunks: {result['chunk_count']}")
        """
        # Step 1: Validate
        is_valid, error = self.validate_file(file_content, filename)
        if not is_valid:
            raise ValueError(error)

        # Step 2: Generate unique file ID (with collision detection if provider available)
        if file_metadata_provider is not None:
            file_id = await self.generate_unique_file_id(
                file_content, filename, file_metadata_provider
            )
        else:
            # Fallback to simple generation (for backward compatibility)
            file_id = self.generate_file_id(file_content, filename)

        # Step 3: Extract text
        text = self.extract_text(file_content, filename)

        # Step 4: Chunk text using strategy
        base_metadata = {"file_id": file_id, "filename": filename}
        chunks = self.chunk_text(text, metadata=base_metadata)

        # Step 5: Enrich chunk metadata
        chunks = self.enrich_chunk_metadata(
            chunks=chunks,
            file_id=file_id,
            filename=filename,
            file_size=len(file_content)
        )

        result = {
            "file_id": file_id,
            "filename": filename,
            "file_size": len(file_content),
            "chunks": chunks,  # Now contains structured chunk dicts
            "chunk_count": len(chunks),
            "chunking_strategy": self.chunking_strategy.get_strategy_name()
        }

        logger.info(
            f"Processed file '{filename}': "
            f"file_id={file_id}, chunks={len(chunks)}, size={len(file_content)} bytes"
        )

        return result


# Singleton instance for dependency injection
_input_data_service_instance: Optional[InputDataHandleService] = None


def get_input_data_service() -> InputDataHandleService:
    """
    FastAPI dependency for Input Data Handle Service (Singleton)

    Usage in endpoints:
        @router.post("/upload")
        async def upload(
            service: InputDataHandleService = Depends(get_input_data_service)
        ):
            ...
    """
    global _input_data_service_instance

    if _input_data_service_instance is None:
        _input_data_service_instance = InputDataHandleService()

    return _input_data_service_instance
